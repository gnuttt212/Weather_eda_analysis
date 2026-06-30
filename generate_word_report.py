import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('BÁO CÁO ĐỒ ÁN: PHÂN TÍCH KHÁM PHÁ DỮ LIỆU (EDA) VÀ DỰ ĐOÁN NHIỆT ĐỘ TP.HCM', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph("Chủ đề: Phân tích Dữ liệu thời tiết & Dự đoán nhiệt độ bằng Machine Learning").style = 'Intense Quote'
    doc.add_paragraph("Vị trí: TP. Hồ Chí Minh, Việt Nam").style = 'Intense Quote'
    doc.add_paragraph("Nguồn dữ liệu: Visual Crossing Weather API").style = 'Intense Quote'
    doc.add_paragraph("Thời gian dữ liệu: 23/06/2023 - 22/06/2026 (1096 ngày ~ 3 năm)").style = 'Intense Quote'
    
    doc.add_heading('1. Mục tiêu đồ án', level=1)
    doc.add_paragraph("Thực hiện quá trình Phân tích Khám phá Dữ liệu (EDA) hoàn chỉnh và áp dụng Machine Learning trên bộ dữ liệu chuỗi thời gian, bao gồm:")
    doc.add_paragraph("Tiền xử lý dữ liệu phức tạp: Chuyển đổi đơn vị từ Imperial (Mỹ) sang Metric (Việt Nam).", style='List Bullet')
    doc.add_paragraph("Trích xuất đặc trưng (Feature Engineering): Lag features, Rolling statistics, phân loại mùa.", style='List Bullet')
    doc.add_paragraph("Thống kê mô tả và trực quan hóa dữ liệu đa chiều (nhiệt độ, lượng mưa, áp suất, tốc độ gió, bức xạ mặt trời, v.v.).", style='List Bullet')
    doc.add_paragraph("Đánh giá và so sánh 5 mô hình Học máy (Machine Learning) để dự đoán nhiệt độ trung bình ngày tiếp theo.", style='List Bullet')
    
    doc.add_heading('2. Tổng quan dữ liệu & Tiền xử lý', level=1)
    doc.add_paragraph("Kích thước dataset: 1096 bản ghi, 33 cột gốc.", style='List Bullet')
    doc.add_paragraph("Tiền xử lý:", style='List Bullet')
    doc.add_paragraph("Loại bỏ các cột không dùng: name, snow, snowdepth, icon, v.v.", style='List Bullet 2')
    doc.add_paragraph("Chuyển đổi đơn vị: Nhiệt độ (°F → °C), Lượng mưa (inches → mm), Gió (mph → km/h), Tầm nhìn (miles → km).", style='List Bullet 2')
    doc.add_paragraph("Điền giá trị thiếu (Missing values imputation) bằng phương pháp nội suy tuyến tính (Linear Interpolation) cho chuỗi thời gian.", style='List Bullet 2')
    doc.add_paragraph("Xử lý ngoại lai (Outliers): Giữ lại toàn bộ các điểm ngoại lai (mưa lớn, gió mạnh) vì đây là các hiện tượng thời tiết cực đoan có ý nghĩa.", style='List Bullet')
    doc.add_paragraph("Feature Engineering:", style='List Bullet')
    doc.add_paragraph("Tạo các biến thời gian: Ngày, Tháng, Năm, Tuần.", style='List Bullet 2')
    doc.add_paragraph("Phân loại mùa: Mùa mưa (T5-T11) và Mùa khô (T12-T4) đặc trưng của TP.HCM.", style='List Bullet 2')
    doc.add_paragraph("Tính toán: Chênh lệch nhiệt độ ngày-đêm, Chênh lệch Điểm sương (Dew Depression), Số giờ ban ngày (Daylight Hours).", style='List Bullet 2')
    doc.add_paragraph("Tạo biến Lag (độ trễ t-1, t-2, t-3) và Rolling (trung bình trượt 3 ngày, 7 ngày).", style='List Bullet 2')

    doc.add_heading('3. Thống kê mô tả & Tương quan', level=1)
    doc.add_paragraph("Nhiệt độ: Trung bình 27.6°C (Biến thiên từ 22.9°C đến 32.9°C).", style='List Bullet')
    doc.add_paragraph("Độ ẩm: Trung bình 78.4%.", style='List Bullet')
    doc.add_paragraph("Lượng mưa: Trung bình 5.4 mm/ngày.", style='List Bullet')
    doc.add_paragraph("Tương quan đáng chú ý:", style='List Bullet')
    doc.add_paragraph("Nhiệt độ & Điểm sương (Dew Point): Tương quan thuận mạnh (r = 0.7581).", style='List Bullet 2')
    doc.add_paragraph("Nhiệt độ & Độ ẩm: Tương quan thuận (r = 0.5460).", style='List Bullet 2')
    doc.add_paragraph("Lượng mưa & Độ ẩm: Tương quan thuận (r = 0.4808).", style='List Bullet 2')
    doc.add_paragraph("Bức xạ mặt trời & Mây che phủ: Tương quan nghịch (r = -0.3815).", style='List Bullet 2')
    
    doc.add_heading('4. Kết quả Trực quan hóa', level=1)
    doc.add_paragraph("Hệ thống đã tự động kết xuất 12 biểu đồ phân tích chuyên sâu (lưu trong thư mục chart/):")
    
    charts = [
        ("bieu_do_1_chuoi_thoi_gian.png", "1. Chuỗi thời gian 4 biến chính (Nhiệt độ, Độ ẩm, Lượng mưa, Gió)."),
        ("bieu_do_2_phan_phoi.png", "2. Histogram phân bố của 8 biến thời tiết khác nhau."),
        ("bieu_do_3_tuong_quan.png", "3. Heatmap ma trận tương quan và Scatter plots."),
        ("bieu_do_4_theo_thang.png", "4. Phân tích nhiệt độ, lượng mưa, độ ẩm theo 12 tháng."),
        ("bieu_do_5_theo_mua.png", "5. Phân tích so sánh Mùa mưa và Mùa khô (Violin plot & Pie chart)."),
        ("bieu_do_6_moving_average.png", "6. Biểu đồ Trung bình trượt 7 và 30 ngày (lọc nhiễu)."),
        ("bieu_do_7_heatmap_thang_nam.png", "7. Ma trận nhiệt độ, độ ẩm theo Tháng x Năm."),
        ("bieu_do_8_pairplot.png", "8. Ma trận biểu đồ cặp phân định theo Mùa."),
        ("bieu_do_9_so_sanh_nam.png", "9. Đồ thị so sánh biến động thời tiết qua các năm."),
        ("bieu_do_10_so_sanh_mo_hinh.png", "10. So sánh trực quan hiệu suất 5 mô hình ML."),
        ("bieu_do_11_feature_importance.png", "11. Mức độ quan trọng của từng Feature từ Random Forest & XGBoost."),
        ("bieu_do_12_ket_qua_mo_hinh.png", "12. Actual vs Predicted của mô hình tốt nhất (kèm Phân phối sai số & QQ-Plot).")
    ]
    
    for filename, caption in charts:
        doc.add_paragraph(caption, style='List Number')
        img_path = os.path.join('chart', filename)
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph(f"[Không tìm thấy ảnh: {filename}]", style='List Bullet 2')
    
    doc.add_heading('5. Mô hình Học máy (Machine Learning)', level=1)
    doc.add_paragraph("Bài toán: Dự đoán nhiệt độ trung bình ngày mai dựa trên 36 đặc trưng thời tiết của ngày hôm nay và các ngày trước đó (Lags & Rolling stats).")
    doc.add_paragraph("Phương pháp chia dữ liệu: Temporal Split (Train: 80%, Test: 20%) — không xáo trộn ngẫu nhiên để tránh Data Leakage trong Time Series.")
    
    doc.add_heading('So sánh 5 Mô hình:', level=2)
    # create table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Mô hình'
    hdr_cells[1].text = 'R² Score (Test)'
    hdr_cells[2].text = 'RMSE (°C)'
    hdr_cells[3].text = 'MAE (°C)'
    hdr_cells[4].text = 'MAPE (%)'
    
    data = [
        ('Lasso Regression (Thắng cuộc)', '0.3504', '1.1667', '0.9199', '3.45'),
        ('Ridge Regression', '0.3401', '1.1759', '0.9234', '3.46'),
        ('Linear Regression', '0.3376', '1.1781', '0.9258', '3.47'),
        ('Random Forest', '0.3241', '1.1900', '0.9562', '3.57'),
        ('XGBoost', '0.3055', '1.2063', '0.9652', '3.60')
    ]
    
    for row in data:
        row_cells = table.add_row().cells
        for i in range(5):
            row_cells[i].text = row[i]
            
    doc.add_heading('Kết luận & Nhận xét', level=2)
    doc.add_paragraph("Thắng cuộc: Mô hình Lasso Regression cho hiệu suất ổn định và tốt nhất trên tập Test, đạt R² = 0.35. Sai số dự đoán trung bình (MAE) khoảng 0.92°C, một mức khá tốt cho dự đoán nhiệt độ trong thực tế.", style='List Bullet')
    doc.add_paragraph("Hiện tượng Overfitting: Random Forest và XGBoost đạt điểm rất cao trên tập Train (R² > 0.93) nhưng lại hoạt động kém trên tập Test. Mô hình dạng cây (Tree-based) gặp khó khăn trong việc dự đoán xu hướng (extrapolation) ngoài khoảng dữ liệu đã học, vốn là rào cản phổ biến khi xử lý Time Series với Temporal Split.", style='List Bullet')
    doc.add_paragraph("Feature Importance: Thông qua Random Forest & XGBoost, các biến quan trọng nhất ảnh hưởng tới nhiệt độ ngày mai bao gồm: Nhiệt độ ngày hôm trước (Temp_Lag1), Chênh lệch nhiệt độ - Điểm sương (DewDepression_C), Áp suất khí quyển (Pressure_mbar) và Nhiệt độ trung bình 3 ngày qua (Temp_Roll3).", style='List Bullet')
    doc.add_paragraph("Cải tiến rõ rệt: So với mô hình đơn giản ban đầu (R² âm), mô hình đa biến có Feature Engineering đã làm chủ bài toán tốt hơn hẳn.", style='List Bullet')
    
    doc.save('BaoCao_EDA_ThoiTiet.docx')
    print("Đã tạo báo cáo BaoCao_EDA_ThoiTiet.docx thành công!")

if __name__ == '__main__':
    main()
